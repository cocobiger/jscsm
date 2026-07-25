# -*- coding: utf-8 -*-
"""生成两份城运对接文档（单一内容源 → 双格式）：
  1) 城运中心对接联调手册（简洁、面向城运研发）
  2) 接口说明文档（完整 API 参考）
  各输出 .docx 与 .md，共 4 个文件。内容严格对齐已部署后端代码。
"""
import os

# ============================================================
# 渲染器：block 列表 → Markdown / Word
# block 类型: ('h1'|'h2'|'h3', text) / ('p', text) / ('bullet', text)
#            / ('code', text) / ('table', headers, rows[, widths])
# ============================================================
def render_md(blocks):
    out = []
    for b in blocks:
        t = b[0]
        if t == 'h1': out.append('# ' + b[1] + '\n')
        elif t == 'h2': out.append('## ' + b[1] + '\n')
        elif t == 'h3': out.append('### ' + b[1] + '\n')
        elif t == 'p': out.append(b[1] + '\n')
        elif t == 'bullet': out.append('- ' + b[1] + '\n')
        elif t == 'code':
            out.append('```\n' + b[1] + '\n```\n')
        elif t == 'table':
            headers, rows = b[1], b[2]
            out.append('| ' + ' | '.join(str(h) for h in headers) + ' |')
            out.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
            for row in rows:
                out.append('| ' + ' | '.join(str(c) for c in row) + ' |')
            out.append('')
    return '\n'.join(out)

def render_docx(blocks, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    CN = '宋体'; HEAD = '黑体'
    def set_cn(run, font=CN):
        run.font.name = font
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font)
    d = Document()
    for b in blocks:
        t = b[0]
        if t == 'h1':
            p = d.add_paragraph(); run = p.add_run(b[1]); run.bold = True
            set_cn(run, HEAD); run.font.size = Pt(16); p.space_after = Pt(6)
        elif t == 'h2':
            p = d.add_paragraph(); run = p.add_run(b[1]); run.bold = True
            set_cn(run, HEAD); run.font.size = Pt(14); p.space_after = Pt(4)
        elif t == 'h3':
            p = d.add_paragraph(); run = p.add_run(b[1]); run.bold = True
            set_cn(run, HEAD); run.font.size = Pt(12); p.space_after = Pt(2)
        elif t == 'p':
            p = d.add_paragraph(); run = p.add_run(b[1]); set_cn(run); run.font.size = Pt(11)
        elif t == 'bullet':
            p = d.add_paragraph(style='List Bullet'); run = p.add_run(b[1])
            set_cn(run); run.font.size = Pt(11)
        elif t == 'code':
            p = d.add_paragraph(); run = p.add_run(b[1])
            run.font.name = 'Consolas'; run.font.size = Pt(9); set_cn(run, 'Consolas')
            p.paragraph_format.left_indent = Inches(0.2)
        elif t == 'table':
            headers, rows = b[1], b[2]
            widths = b[3] if len(b) > 3 else None
            tb = d.add_table(rows=1, cols=len(headers))
            tb.style = 'Table Grid'; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            hc = tb.rows[0].cells
            for i, h in enumerate(headers):
                hc[i].text = ''
                run = hc[i].paragraphs[0].add_run(str(h)); run.bold = True
                set_cn(run, HEAD); run.font.size = Pt(10)
            for row in rows:
                cells = tb.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ''
                    run = cells[i].paragraphs[0].add_run(str(val))
                    set_cn(run); run.font.size = Pt(9.5)
            if widths:
                for i, w in enumerate(widths):
                    for r in tb.rows:
                        r.cells[i].width = Inches(w)
    d.save(path)

# ============================================================
# 内容：标题 / 说明文案（单一来源）
# ============================================================
BASE = 'http://[驾驶舱对外地址]'

# 通用事件字段表（入站）
EVENT_FIELDS = (
    ['字段', '必填', '类型', '说明'],
    [
        ['eventId', '是', 'string', '事件ID（规则 sjzl-摄像头id-毫秒时间戳）；作为 warning.id 幂等键，重复推送不重复落库'],
        ['eventTime', '是', 'string', '事件时间（上海本地时，如 2024-09-15 07:14:25；无时区标记，落库补 +08:00）'],
        ['cameraId', '是', 'string', '摄像头ID → 关联 iot_channels 通道与通道名'],
        ['eventType', '是', 'int', '枚举 1~17 → aiType（映射见下文）'],
        ['subType', '否', 'string', '子类型'],
        ['confirm', '否', 'int', '1 已审核 / 0 未审核 → 决定告警级别（未审核=1，已审核=2）'],
        ['eventImgSmall', '否', 'string', '小图网络地址 → picUrl（经 /api/iot-image 代理预览）'],
        ['eventImgBig', '否', 'string', '大图网络地址（优先于 small）'],
        ['longitude', '否', 'number', '经度；落库口径：经度 = longitude'],
        ['latitude', '否', 'number', '纬度；落库口径：纬度 = latitude（已与城运确认）'],
        ['districtName', '否', 'string', '区县 → location'],
        ['townName', '否', 'string', '乡镇/街道 → location'],
        ['address', '否', 'string', '详细地址 → location'],
        ['elevation', '否', 'number', '云台上下角度'],
        ['azimuth', '否', 'number', '云台水平角度'],
        ['absoluteZoom', '否', 'number', '放大倍数'],
        ['processEventId', '否', 'string', '流程事件标识'],
        ['processEventStatus', '否', 'int', '流程事件状态：0 未识别 / 1 识别到（可用于流程结束判断）'],
        ['watermarkImage', '否', 'string', '水印图地址'],
        ['count', '否', 'int', '识别物数量'],
        ['total', '否', 'int', '总数'],
        ['presetPosNum', '否', 'string', '预置点编号'],
        ['presetPosName', '否', 'string', '预置点名称'],
        ['distance', '否', 'number', '预估距离'],
    ],
)

# eventType → aiType 映射
EVENT_TYPE_MAP = (
    ['平台 eventType', '含义', '驾驶舱 aiType'],
    [
        ['1', '工程车作业', '工程车作业'],
        ['2', '工程车数量', '工程车数量'],
        ['3', '烟尘', '烟尘'],
        ['4', '工地裸露地未覆盖', '堆头未覆盖（已确认对应）'],
        ['5', '生物质燃烧', '生物质燃烧'],
        ['6', '烟囱烟雾', '烟囱烟雾'],
        ['7', '扬尘', '扬尘'],
        ['8', '人员入侵', '人员入侵'],
        ['9', '卡车脏车', '卡车脏车'],
        ['10', '脏车', '脏车'],
        ['11', '车辆遗撒', '车辆遗撒'],
        ['12', '建渣未覆盖', '建渣未覆盖'],
        ['16', '车辆冒装', '车辆冒装'],
        ['17', '工业烟羽', '工业烟羽'],
    ],
)

# 回调请求字段表
CALLBACK_FIELDS = (
    ['字段', '必填', '说明'],
    [
        ['disposal_status', '是', '处置状态：processing（受理中）/ closed（已结案）；其他取值默认按 processing 处理'],
        ['disposal_result', '否', '处置结论 / 结果文本'],
        ['disposal_operator', '否', '处置人 / 责任单位'],
        ['disposal_time', '否', '处置完成时间（建议上海时字符串）'],
        ['(兼容别名)', '否', 'status / result / remark / operator / disposalResult / disposalOperator / disposalTime，字段名可由 callback_field_mapping 适配'],
    ],
)

# 错误码表
ERROR_CODES = (
    ['HTTP', 'code', 'message', '含义 / 处理'],
    [
        ['200', '200', '请求已成功', '成功；data 含业务字段'],
        ['400', '400', '缺少 push_id / 缺少 eventId', '缺少必填关联键；请携带 X-Push-Id 头或报文体 push_id'],
        ['401', '401', '令牌无效', '配置了令牌但 X-Callback-Token 不符；请核对令牌'],
        ['403', '403', '来源 IP 不在白名单', '调用方 IP 未在 CHENGYUN_ALLOW_IPS；联系驾驶舱侧加白'],
        ['404', '404', '推送记录不存在', 'push_id 在驾驶舱无对应推送记录；确认 push_id 正确'],
        ['500', '500', '处理失败', '服务端异常；查看驾驶舱日志'],
    ],
)

# 部署配置表
DEPLOY_VARS = (
    ['变量', '默认值', '说明'],
    [
        ['CHENGYUN_ALLOW_IPS', '127.0.0.1', '入站/回调允许的客户端 IP 白名单（逗号分隔）；生产须设为城运平台出口 IP 与驾驶舱自身 IP'],
        ['CHENGYUN_CALLBACK_TOKEN', '（空）', '可选。若设置，入站/回调须带 X-Callback-Token 头且相等，否则 401；当前不强制（令牌预留）'],
        ['CHENGYUN_IMG_HOSTS', '10.120.49.14', '城运平台图片/短视频域名白名单，供 /api/iot-image 代理放行'],
        ['SMART_PUSH_CALLBACK_URL', '（空）', '可选。驾驶舱对外回调地址，注入推送报文 callback_url 变量，便于城运回传'],
    ],
)

# ============================================================
# 文档 1：对接联调手册（简洁）
# ============================================================
def build_manual():
    B = []
    B.append(('h1', '城运中心对接联调手册'))
    B.append(('p', '系统：环保局驾驶舱系统（JSC）。本文档供城运中心（全景影像视频平台 / 三级城运中心）研发对接驾驶舱使用，描述两条链路的接口与联调步骤。所有响应体统一为 {"code":int,"message":str,"data":{...}}。'))
    B.append(('h2', '一、对接链路总览'))
    B.append(('bullet', '链路 A（入站 · 城运 → 驾驶舱）：全景影像视频平台按《全景影像视频平台接口规范 V1.0》调用我方订阅接口，推送摄像头识别事件与短视频。这是驾驶舱 AI 视频告警（堆头未覆盖、扬尘等）的标准事件来源。'))
    B.append(('bullet', '链路 B（出站闭环 · 驾驶舱 → 城运 → 驾驶舱）：预警经城运中心 API 推送至政务系统处置，处置完成后城运中心回调我方接口回传处置结论，形成 pushed → processing → closed 闭环。'))
    B.append(('h2', '二、网络与鉴权'))
    B.append(('bullet', '接口部署于电子政务网隔离区；入站与回调接口采用「来源 IP 白名单 + 可选令牌」鉴权（复用同一守卫 chengyunGuard）。'))
    B.append(('bullet', 'IP 白名单：由驾驶舱侧在 CHENGYUN_ALLOW_IPS 配置城运平台出口 IP；不在白名单返回 403。'))
    B.append(('bullet', '令牌（可选）：若驾驶舱配置了 CHENGYUN_CALLBACK_TOKEN，请求须带请求头 X-Callback-Token 且与之相等，否则 401。当前阶段令牌不强制（预留）。'))
    B.append(('bullet', '/client/* 不经过 /api 会话鉴权；/api/smart-push/callback 已加入公开路径绕过会话鉴权。'))
    B.append(('h2', '三、接口速查表'))
    B.append(('table',
        ['方向', '方法', '路径', '用途', '鉴权'],
        [
            ['入站', 'POST', '/client/handle_event', '推送摄像头识别事件', 'IP 白名单(+令牌)'],
            ['入站', 'POST', '/client/handle_event_other', '推送关联短视频', 'IP 白名单(+令牌)'],
            ['出站(城运接收)', 'POST', '<城运侧提供的接收地址 api_url>', '驾驶舱推送预警事件', '由城运侧定义'],
            ['回调', 'POST', '/api/smart-push/callback', '城运回传处置结论', 'IP 白名单(+令牌)'],
            ['(内部)', 'POST', '/api/smart-push/history/:id/close', '人工一键结案', '会话 Bearer'],
        ],
        [0.9, 0.7, 2.6, 1.8, 1.4]))
    B.append(('h2', '四、链路 A：入站接口（城运 → 驾驶舱）'))
    B.append(('h3', '4.1 POST /client/handle_event'))
    B.append(('p', '请求体为 JSON，字段如下（与《全景影像视频平台接口规范 V1.0》一致）：'))
    B.append(('table', EVENT_FIELDS[0], EVENT_FIELDS[1], [1.7, 0.5, 0.8, 3.5]))
    B.append(('p', '请求示例：'))
    B.append(('code',
        'curl -X POST ' + BASE + '/client/handle_event \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -H "X-Callback-Token: <若启用>" \\\n'
        '  -d \'{\n'
        '    "eventId": "sjzl-119-202409150714251",\n'
        '    "eventTime": "2024-09-15 07:14:25",\n'
        '    "cameraId": "119",\n'
        '    "eventType": 4,\n'
        '    "confirm": 1,\n'
        '    "eventImgSmall": "http://10.120.49.14:30100/img/s.jpg",\n'
        '    "eventImgBig": "http://10.120.49.14:30100/img/b.jpg",\n'
        '    "longitude": 108.3893,\n'
        '    "latitude": 30.8050,\n'
        '    "districtName": "万州区",\n'
        '    "townName": "某街道",\n'
        '    "address": "某路某号"\n'
        '  }\''))
    B.append(('p', '响应示例（成功）：'))
    B.append(('code', '{"code":200,"message":"请求已成功","data":{}}'))
    B.append(('bullet', 'eventId 幂等：以 eventId 为 warning.id 做 upsert，重复推送不重复落库、保留首见时间。'))
    B.append(('bullet', '经纬度口径（已确认）：经度 = longitude 字段，纬度 = latitude 字段。'))
    B.append(('bullet', 'eventType 枚举映射（节选，完整见接口说明文档）：4 → 堆头未覆盖；7 → 扬尘；3 → 烟尘；5 → 生物质燃烧。'))
    B.append(('h3', '4.2 POST /client/handle_event_other（短视频）'))
    B.append(('p', '请求体：eventIds（事件ID，可逗号分隔）、fileUrl（录像文件地址）、cameraId。'))
    B.append(('code',
        'curl -X POST ' + BASE + '/client/handle_event_other \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"eventIds":"sjzl-119-202409150714251","fileUrl":"http://10.120.49.14:30100/vid/xxx.mp4","cameraId":"119"}\''))
    B.append(('p', '响应示例：'))
    B.append(('code', '{"code":200,"message":"请求已成功","data":{"updated":1}}'))
    B.append(('h2', '五、链路 B：出站闭环'))
    B.append(('h3', '5.1 驾驶舱 → 城运 推送（城运侧接收）'))
    B.append(('bullet', '城运侧需提供接收地址（api_url），驾驶舱按预案模板 POST 推送预警事件。'))
    B.append(('bullet', '推送报文（城运侧会收到的字段）：event_type / location / lat / lon / level / value / standard / description / time / trigger_count / event_ids，以及关联键 push_id（=本次推送记录 ID）、callback_url（可选，驾驶舱对外回调地址）。'))
    B.append(('bullet', '回执关联：驾驶舱推送的 HTTP 响应头带 X-Push-Id: <push_id>；同时报文体模板变量含 push_id。城运侧请保存该 push_id，回调时原样回传。'))
    B.append(('p', '城运侧收到的推送报文示例（取决于预案模板，字段名可配）：'))
    B.append(('code',
        'POST <城运接收地址>  Content-Type: application/json\n'
        '{\n'
        '  "event_type": "堆头未覆盖",\n'
        '  "location": "九龙沙场",\n'
        '  "lon": 108.38, "lat": 30.80, "level": 2,\n'
        '  "trigger_count": 23, "event_ids": "id1,id2",\n'
        '  "push_id": "9f3a...uuid",\n'
        '  "callback_url": "' + BASE + '/api/smart-push/callback"\n'
        '}'))
    B.append(('h3', '5.2 城运 → 驾驶舱 回调 POST /api/smart-push/callback'))
    B.append(('bullet', '关联键（三选一，优先级从高到低）：请求头 X-Push-Id > 报文体 push_id > 报文体 event_id；三者皆缺返回 400。'))
    B.append(('p', '请求体字段：'))
    B.append(('table', CALLBACK_FIELDS[0], CALLBACK_FIELDS[1], [1.8, 0.5, 4.2]))
    B.append(('p', '请求示例：'))
    B.append(('code',
        'curl -X POST ' + BASE + '/api/smart-push/callback \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -H "X-Callback-Token: <若启用>" \\\n'
        '  -H "X-Push-Id: 9f3a...uuid" \\\n'
        '  -d \'{\n'
        '    "disposal_status": "closed",\n'
        '    "disposal_result": "已现场核查并责令整改苫盖",\n'
        '    "disposal_operator": "城运中心-张三",\n'
        '    "disposal_time": "2024-09-15 09:30:00"\n'
        '  }\''))
    B.append(('p', '响应示例（成功）：'))
    B.append(('code', '{"code":200,"message":"请求已成功","data":{"status":"closed"}}'))
    B.append(('bullet', '状态机：disposal_status=processing → 关联事件置 processing（可回退）；disposal_status=closed → 关联事件置 closed（不可逆），并写回处置结论、标记待结案。'))
    B.append(('bullet', '超时：事件处于 pushed 超过 24h 无回调，驾驶舱前端标红提示人工跟办。'))
    B.append(('h2', '六、联调检查清单'))
    B.append(('bullet', '驾驶舱侧已将城运出口 IP 加入 CHENGYUN_ALLOW_IPS 并重启服务。'))
    B.append(('bullet', '用已知 eventId 推送一次，确认返回 200 且驾驶舱落库（source=chengyun-platform）。'))
    B.append(('bullet', '重复推送同一 eventId，确认幂等不重复落库。'))
    B.append(('bullet', '驾驶舱侧配置测试推送预案推送到城运接收地址，确认城运收到 push_id（响应头 X-Push-Id）。'))
    B.append(('bullet', '城运用收到的 push_id 调 /api/smart-push/callback（processing），确认状态流转。'))
    B.append(('bullet', '再调 closed，确认事件置 closed 且处置写回。'))
    B.append(('bullet', '缺 push_id 调用，确认返回 400；非白名单 IP 调用，确认返回 403。'))
    B.append(('h2', '七、错误码'))
    B.append(('table', ERROR_CODES[0], ERROR_CODES[1], [0.7, 0.7, 2.0, 3.1]))
    B.append(('h2', '八、待与城运确认'))
    B.append(('bullet', '驾驶舱对外回调地址（公网/政务网域名）— 影响城运侧调用 /api/smart-push/callback 的 URL。'))
    B.append(('bullet', '城运侧接收地址 api_url（驾驶舱推送目标）。'))
    B.append(('bullet', '是否启用令牌鉴权（CHENGYUN_CALLBACK_TOKEN），以及令牌交换方式。'))
    return B

# ============================================================
# 文档 2：接口说明文档（完整）
# ============================================================
def build_api():
    B = []
    B.append(('h1', '驾驶舱系统（JSC）接口说明文档 · 城运中心对接'))
    B.append(('p', '本文档为驾驶舱系统（JSC）与城运中心（全景影像视频平台 / 三级城运中心）对接的完整 API 参考，与已部署后端代码严格一致。所有响应体统一为 {"code":int,"message":str,"data":{...}}。'))
    B.append(('h2', '一、通用约定'))
    B.append(('bullet', '协议：HTTP + JSON（application/json）。'))
    B.append(('bullet', '时区：所有时间字段为上海时（Asia/Shanghai），格式 YYYY-MM-DD HH:MM:SS。'))
    B.append(('bullet', '鉴权：入站与回调接口使用 chengyunGuard（来源 IP 白名单 CHENGYUN_ALLOW_IPS + 可选令牌 X-Callback-Token）。'))
    B.append(('bullet', '统一响应结构：{"code":<int>,"message":<str>,"data":{...}}。'))
    B.append(('h2', '二、术语'))
    B.append(('bullet', 'push_id：驾驶舱 smart_push_history 记录 ID，作为城运回调的关联键。'))
    B.append(('bullet', 'eventId：城运平台事件 ID，作为驾驶舱 warnings 的幂等键（warning.id）。'))
    B.append(('h2', '三、入站接口（城运 → 驾驶舱）'))
    B.append(('h3', '3.1 POST /client/handle_event'))
    B.append(('p', '方法/路径：POST /client/handle_event（经 nginx /client/ 转发至后端 7170）。鉴权：chengyunGuard（IP 白名单 + 可选 X-Callback-Token）。'))
    B.append(('p', '请求参数：'))
    B.append(('table', EVENT_FIELDS[0], EVENT_FIELDS[1], [1.5, 0.45, 0.7, 3.85]))
    B.append(('p', '落库说明：以 eventId 为 warning.id 做幂等 upsert（INSERT OR REPLACE），重复推送不重复落库、保留首见时间与既有处置状态；source=chengyun-platform；经纬度口径 经度=longitude / 纬度=latitude；confirm=1 则 level=2，否则 level=1。'))
    B.append(('p', '响应：'))
    B.append(('bullet', '成功：200 {"code":200,"message":"请求已成功","data":{}}'))
    B.append(('bullet', '缺 eventId：400 {"code":400,"message":"缺少 eventId","data":{}}'))
    B.append(('bullet', '非白名单 IP：403 {"code":403,"message":"来源 IP 不在白名单","data":{}}'))
    B.append(('bullet', '令牌不符：401 {"code":401,"message":"令牌无效","data":{}}'))
    B.append(('h3', '3.2 POST /client/handle_event_other'))
    B.append(('p', '方法/路径：POST /client/handle_event_other；鉴权同 3.1。'))
    B.append(('p', '请求参数：eventIds（事件ID，可逗号分隔）、fileUrl（录像文件地址）、cameraId。'))
    B.append(('p', '响应：200 {"code":200,"message":"请求已成功","data":{"updated":N}}（N 为成功关联的事件数）。行为：将 fileUrl 关联到对应 warning（按 eventIds 匹配），作为结案报告证据区视频证据。'))
    B.append(('h2', '四、出站接口（驾驶舱 → 城运 → 驾驶舱）'))
    B.append(('h3', '4.1 驾驶舱 → 城运 推送（城运侧接收）'))
    B.append(('bullet', '触发：POST /api/smart-push/events 按 smart_push_rules 匹配 → 调用预案 api_url 推送；推送成功关联事件置 pushed、历史置 pushed。'))
    B.append(('bullet', '推送报文模板变量（注入 body_template）：标准变量 event_type / location / lat / lon / level / value / standard / description / time / trigger_count / event_ids + 关联变量 push_id（本次 smart_push_history 记录 ID）+ callback_url（可选，配置 SMART_PUSH_CALLBACK_URL 时注入）。'))
    B.append(('bullet', '回执关联头：推送 HTTP 响应携带响应头 X-Push-Id: <push_id>，城运中心用于回调时关联本次推送。'))
    B.append(('h3', '4.2 POST /api/smart-push/callback（处置反馈回调 · 核心闭环）'))
    B.append(('p', '方法/路径：POST /api/smart-push/callback；鉴权：复用 chengyunGuard（IP 白名单 + 可选 X-Callback-Token）。'))
    B.append(('bullet', '关联键（三选一，优先级从高到低）：请求头 X-Push-Id > 报文体 push_id > 报文体 event_id；三者皆缺返回 400。'))
    B.append(('p', '请求参数：'))
    B.append(('table', CALLBACK_FIELDS[0], CALLBACK_FIELDS[1], [1.6, 0.5, 4.4]))
    B.append(('p', '响应：'))
    B.append(('bullet', '成功：200 {"code":200,"message":"请求已成功","data":{"status":"<processing|closed>"}}'))
    B.append(('bullet', '推送记录不存在：404 {"code":404,"message":"推送记录不存在","data":{}}'))
    B.append(('bullet', '令牌不符：401；非白名单 IP：403；缺 push_id：400。'))
    B.append(('p', '状态机（smart_push_events）：收到 processing → 关联事件置 processing（可回退）；收到 closed → 关联事件置 closed（不可逆，不得回退为 processing）+ 推送历史写回处置结论（disposal_result / disposal_operator / callback_time=反馈时间），标记待结案。'))
    B.append(('h3', '4.3 POST /api/smart-push/history/:id/close（人工一键结案）'))
    B.append(('p', '方法/路径：POST /api/smart-push/history/:id/close；鉴权：会话鉴权（Bearer），操作人取当前登录用户。'))
    B.append(('bullet', '响应：成功 {"ok":true,"status":"closed"}；记录不存在 {"ok":false,"error":"推送记录不存在"}（HTTP 404）。'))
    B.append(('bullet', '行为：历史置 closed + 关联事件置 closed + 写回处置人 / 结案时间（disposal_result 默认"驾驶舱人工结案"）。'))
    B.append(('h3', '4.4 GET /api/smart-push/history'))
    B.append(('p', '查询推送历史。参数：status（pushed / processing / closed / timeout 特殊：内存过滤 is_timeout=1）、event_type（按事件类型）、limit（默认 100）。返回记录含 is_timeout 字段（status=pushed 且 24h 内无回调为 1）。'))
    B.append(('h2', '五、数据模型与枚举'))
    B.append(('h3', '5.1 平台 eventType → 驾驶舱 aiType 映射'))
    B.append(('table', EVENT_TYPE_MAP[0], EVENT_TYPE_MAP[1], [1.4, 2.2, 2.9]))
    B.append(('h3', '5.2 事件状态机'))
    B.append(('table',
        ['状态', '含义', '触发方'],
        [
            ['pending', '预警产生，待处理', '系统'],
            ['pushed', '已推送城运中心，等待政务受理', '驾驶舱推送成功'],
            ['processing', '政务系统已受理、处置中', '城运回调'],
            ['closed', '政务处置完成，事件待结案', '城运回调'],
            ['handled', '本地手动处置完成（未推送城运的事件）', '值守人员'],
        ],
        [1.4, 3.0, 2.1]))
    B.append(('h3', '5.3 经纬度落库口径（已确认）'))
    B.append(('bullet', '经度 = 平台 longitude 字段；纬度 = 平台 latitude 字段。'))
    B.append(('h2', '六、错误码'))
    B.append(('table', ERROR_CODES[0], ERROR_CODES[1], [0.7, 0.7, 2.0, 3.1]))
    B.append(('h2', '七、部署配置（环境变量，经 systemd EnvironmentFile 注入）'))
    B.append(('p', '配置位于 /opt/jsc/backend/iotcloud.env，改完 systemctl restart jsc-backend 生效。'))
    B.append(('table', DEPLOY_VARS[0], DEPLOY_VARS[1], [2.2, 1.1, 3.3]))
    B.append(('h2', '八、变更记录'))
    B.append(('bullet', 'V1.0（2026-07-11）：初版接口说明，含入站事件/短视频、出站推送与处置反馈回调闭环、人工结案、查询、枚举映射、错误码与部署配置。'))
    return B

if __name__ == '__main__':
    out_dir = r'E:\CC work\CC jsc\docs'
    os.makedirs(out_dir, exist_ok=True)
    manual = build_manual()
    api = build_api()
    # Word
    render_docx(manual, os.path.join(out_dir, '城运中心对接联调手册.docx'))
    render_docx(api, os.path.join(out_dir, 'JSC接口说明文档_城运对接.docx'))
    # Markdown
    with open(os.path.join(out_dir, '城运中心对接联调手册.md'), 'w', encoding='utf-8') as f:
        f.write(render_md(manual))
    with open(os.path.join(out_dir, 'JSC接口说明文档_城运对接.md'), 'w', encoding='utf-8') as f:
        f.write(render_md(api))
    print('DONE: 4 files generated in', out_dir)
