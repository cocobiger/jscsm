# -*- coding: utf-8 -*-
"""生成两份 enriched 文档：在原有内容基础上新增「事件处置归档与结案PDF」模块，
V1.2 重点：第九章接入「城运中心对接闭环」（预警→推送城运→政务处置→反馈回舱→结案归档）。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CN_FONT = '宋体'
HEADER_FONT = '黑体'

def set_cn(run, font=CN_FONT):
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font)

def add_heading(doc, text, level=1, size=None, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    set_cn(run, HEADER_FONT)
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(15 if level == 1 else 13 if level == 2 else 11.5)
    if color:
        run.font.color.rgb = color
    p.space_after = Pt(6)
    return p

def add_para(doc, text, bold=False, size=11, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    set_cn(run)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_cn(run)
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        set_cn(run, HEADER_FONT)
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn(run)
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t

# ============================================================
# 文档 1：需求沟通说明（会议纪要）
# ============================================================
def build_doc1(path):
    d = Document()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('环保局“驾驶舱”系统建设需求沟通说明'); r.bold = True; set_cn(r, HEADER_FONT); r.font.size = Pt(18)
    p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('——2026年7月10日项目沟通会议纪要及待确认事项（含事件处置归档与结案报告模块）'); set_cn(r2); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x55,0x55,0x55)

    add_heading(d, '会议背景', 1)
    add_para(d, '2026年7月10日，我方与环保局召开“驾驶舱”系统建设项目沟通会议，就系统建设目标、展示内容与落地条件进行了初步交流。根据会议讨论情况，现梳理出三项需与环保局进一步沟通并予确认的重点事项，以便我方推进后续硬件准备与需求文档编制工作。')

    add_heading(d, '需与环保局确认的事项', 1)

    add_heading(d, '（一）驾驶舱适配终端的硬件配置要求', 2)
    add_para(d, '为保障驾驶舱系统（尤其是大屏集中展示与交互演示环节）稳定、流畅运行，建议部署终端满足以下最低硬件配置标准：')
    add_table(d,
        ['配置项','最低要求','说明'],
        [['内存','≥ 16GB','建议 32GB 以预留余量'],
         ['CPU','≥ Intel i3-12400','或同级别及以上型号'],
         ['显卡','≥ NVIDIA RTX 2060','需支持多路视频硬解码'],
         ['硬盘','固态硬盘（SSD）','系统及数据盘均建议 SSD'],
         ['显示器','分辨率 ≥ 2560×1440','2K 及以上，适配大屏投屏'],
         ['网络','电子政务网独享带宽 ≥ 20MB','可在网关侧进行带宽分配']],
        widths=[1.4, 2.2, 3.0])
    add_para(d, '配置依据：驾驶舱需同时承载约 30 路视频信号在大屏集中展示；若用于向上级领导汇报，还将涉及气象模拟及实时数据回传至电子地图的实景演示，几乎全部为实时（“活”）数据，且交互操作演示的功能点较多。因此终端需具备较强的计算、图形处理与网络传输能力，以保障多路视频、实时地图与交互演示的整体效果稳定流畅。')

    add_heading(d, '（二）基于重庆两江新区市驾驶舱的改造对接需求', 2)
    add_para(d, '驾驶舱将参照“重庆两江新区市驾驶舱”的成熟模式进行改造对接。为明确改造范围、功能边界与数据对接方案，我方将编制一份图文并茂的需求文档，作为后续系统对接与开发的重要依据。该文档拟包含以下内容：')
    add_bullet(d, '界面参考：提供重庆两江新区市驾驶舱的界面截图，并对各功能点辅以文字说明，明确驾驶舱的页面布局、模块划分与交互方式。')
    add_bullet(d, '展示数据清单：逐项列明驾驶舱需展示的数据指标（如空气质量、污染源监控、视频点位等），并说明其展示形式（图表、地图、列表、视频等）。')
    add_bullet(d, '数据来源说明：逐一标注上述数据的来源，包括对接的业务系统、数据接口（API/数据库）、以及是否需人工填报或外部导入，明确数据获取路径与责任方。')
    add_bullet(d, '表格导入导出能力：明确各数据是否支持通过表格（Excel/CSV）进行批量导入与导出，以及导入导出的字段范围、频率与操作权限，便于日常运维与离线补录。')
    add_para(d, '恳请环保局对上述文档范围内的数据范围、来源系统及表格导入导出方式予以确认或补充，以便我方据此形成正式需求规格并启动对接开发。')

    add_heading(d, '（三）事件处置归档与结案报告模块（新增，待确认）', 2)
    add_para(d, '为落实环境执法“闭环留痕”要求，驾驶舱拟新增「事件处置 → 归档 → 结案PDF」能力，并接入**三级城运中心**实现事件自动流转处置的闭环。设计要点如下：')
    add_bullet(d, '闭环链路：预警产生 → 驾驶舱经城运中心 API 推送事件 → 三级城运中心流转至政务系统处置 → 处置完成后城运中心将处置结果**反馈回驾驶舱** → 自动更新事件状态并写回处置信息 → 形成“待结案” → 一键生成 PDF 结案归档。')
    add_bullet(d, '事件范围：覆盖驾驶舱全部预警。命中城运推送规则的经城运中心流转处置；未命中的仍可由值守人员本地处置（handled），两类事件均可生成结案报告。')
    add_bullet(d, '处置留痕：城运回调自动写回「处置措施、处置结果、责任单位、责任人、现场照片、政务派单号、处置/反馈时间」；本地处置由值守人员录入同样字段。')
    add_bullet(d, '结案报告版式：参照万州区生态环境局《气体溯源系统》已生成的“无人机大气巡查污染溯源报告”版式（事件概况 → 证据与监测分析 → 气象条件 → 溯源分析 → 处置记录 → 综合研判结论），确保执法文书风格统一。')
    add_bullet(d, '溯源章节条件填充：报告中的“监测分析 / 气象条件 / 溯源分析”等依赖气体溯源数据的章节，按事件类型条件填充——关联到无人机气体巡查任务的事件填入真实溯源数据；纯 IoT 视频 / 空气质量事件标注“无溯源数据 / 不适用”，不强行编造。')
    add_bullet(d, '生成方式：报告在前端浏览器端由 html2canvas + jsPDF 生成并下载，与气体溯源系统技术路线一致，后端改动小。')
    add_para(d, '需环保局重点确认的要点：', bold=True)
    add_bullet(d, '处置留痕字段是否满足执法存档要求（是否需扩展字段、是否需电子签章）。')
    add_bullet(d, '结案报告是否需纳入统一档案管理系统 / 纸质打印归档 / 审计留存期限。')
    add_bullet(d, '对无溯源数据的事件，报告“溯源”章节标注“不适用”是否被接受。')
    add_bullet(d, '处置与归档的操作权限（值守人员可处置/生成、管理员可配置对接与归档导出）如何划分。')
    add_bullet(d, '城运中心对接要点：事件关联键（event_id）约定、回调接口字段规范、回调鉴权令牌/来源 IP 白名单的交换方式、推送重试与反馈超时阈值等，需与城运中心（三级城运）确认接口能力。')

    add_heading(d, '下一步建议', 1)
    add_bullet(d, '请环保局就上述第（一）项硬件配置标准予以确认，并据此推进终端采购与环境准备。')
    add_bullet(d, '请环保局就第（二）项中的数据范围、来源及导入导出方式提供意见，我方将据此完成图文需求文档编制并提交评审。')
    add_bullet(d, '请环保局就第（三）项事件处置归档与结案报告模块（含城运中心对接闭环）的字段、签章与归档要求予以确认，我方将据此补充需求规格说明书对应章节。')
    add_bullet(d, '建议双方约定后续沟通时间，对三项事项形成书面确认，作为项目推进的基准；并协调城运中心（三级城运）明确事件推送与处置反馈接口规范。')

    add_para(d, '报送：环保局', size=11)
    add_para(d, '日期：2026年7月11日', size=11)

    d.save(path)
    print('saved doc1 ->', path)

# ============================================================
# 文档 2：需求规格说明书（模板 + 新增模块实填）
# ============================================================
def build_doc2(path):
    d = Document()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('环保局“驾驶舱”系统需求规格说明书'); r.bold = True; set_cn(r, HEADER_FONT); r.font.size = Pt(18)
    p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('（参照重庆两江新区市驾驶舱改造对接 · 含事件处置归档与结案报告模块 · V1.6 含城运中心对接闭环 + 堆头未覆盖事件报告示例 + 城运视频平台事件接入 + 入站接口已落地 + 含接口契约附录（实现版））'); set_cn(r2); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x55,0x55,0x55)

    add_heading(d, '文档信息', 1)
    add_table(d,
        ['文档名称','版本','编制单位','编制日期','评审单位','文档状态'],
        [['环保局“驾驶舱”系统需求规格说明书','V1.6（含城运中心对接闭环 + 堆头未覆盖事件报告示例 + 城运视频平台事件接入 + 入站接口已落地）','【请填写】','2026年7月11日','环保局','待评审 / 草稿']],
        widths=[2.4,1.6,1.2,1.2,1.0,1.2])

    add_heading(d, '文档目的与范围', 1)
    add_para(d, '本文档用于明确环保局“驾驶舱”系统的建设目标、参考原型、功能模块、展示数据、数据来源及表格导入导出需求，作为系统对接与开发实施的基准依据。驾驶舱将参照“重庆两江新区市驾驶舱”的成熟模式进行改造对接。')
    add_para(d, '适用范围：驾驶舱大屏展示、交互演示、数据对接、事件处置与结案归档（含三级城运中心对接闭环）及配套运维能力。不包括硬件采购（详见《需求沟通说明》第一项）与底层网络基础设施建设。')
    add_para(d, '本次 V1.6 在 V1.5 基础上，将第九章 9.7 城运视频平台事件接入（/client/handle_event、/client/handle_event_other）落地为可运行后端代码：新增 upsertWarningFromChengyun 幂等落库（eventType→aiType 映射、cameraId→通道、经纬度以经度=Longitude/纬度=Latitude 落库、source=chengyun-platform）、nginx /client/ 转发、IP 白名单 + 令牌预留守卫、/api/iot-image 代理域名白名单扩展；并修复落库时间时区（+08:00）与首见时间保留。9.7.5 补充三项部署环境变量（CHENGYUN_ALLOW_IPS / CHENGYUN_CALLBACK_TOKEN / CHENGYUN_IMG_HOSTS）。')

    add_heading(d, '参考原型：重庆两江新区市驾驶舱', 1)
    add_para(d, '[截图占位] 重庆两江新区市驾驶舱——总览/首页截图')
    add_para(d, '[截图占位] 重庆两江新区市驾驶舱——重点模块截图（如监测、视频、地图等）')
    add_para(d, '注：原型功能点逐项比对详见下一节，功能设计需与环保局确认后定稿。')

    add_heading(d, '原型功能点说明', 1)
    add_table(d,
        ['模块','功能点','功能说明','原型参考'],
        [['【模块名】','【功能点】','【描述该功能点的作用与交互】','见 2.1 截图'],
         ['【模块名】','【功能点】','【描述】','见 2.1 截图'],
         ['【模块名】','【功能点】','【描述】','见 2.1 截图']],
        widths=[1.2,1.5,3.0,1.0])

    add_heading(d, '驾驶舱功能模块与页面布局', 1)
    add_para(d, '[截图占位] 本驾驶舱页面布局示意图（标注各区域对应的模块）')
    add_heading(d, '功能模块清单', 2)
    add_table(d,
        ['序号','模块名称','功能描述','对应原型功能点','优先级','备注'],
        [['1','【模块】','【描述】','见 2.2','高',''],
         ['2','【模块】','【描述】','见 2.2','中',''],
         ['3','【模块】','【描述】','见 2.2','低',''],
         ['4','事件处置与结案归档（含城运对接）','对预警经城运中心 API 自动推送、接收政务处置反馈并写回处置信息，形成闭环；支持本地处置兜底；一键生成参照气体溯源报告版式的 PDF 结案归档报告；溯源章节按事件类型条件填充。','新增（见第九章）','高','执法闭环留痕']],
        widths=[0.5,1.6,3.0,1.0,0.6,1.0])

    add_heading(d, '展示数据清单', 1)
    add_table(d,
        ['序号','数据指标','展示形式','更新频率','备注'],
        [['1','【如：空气质量指数 AQI】','【图表/地图/列表/视频】','【实时/5分钟】',''],
         ['2','【如：30路视频点位】','视频','实时',''],
         ['3','【如：污染源监控】','【地图/列表】','【频率】',''],
         ['4','事件处置状态流（pending/pushed/processing/closed/handled）','列表 + 状态标签','实时（推送/回调驱动）','来源 warnings + 城运回调'],
         ['5','城运推送与处置反馈记录','列表 + 详情抽屉','实时（推送/回调时写入）','来源 smart_push_history + disposition'],
         ['6','结案归档报告（PDF）','PDF 文档（前端生成下载）','按需生成','版式参照气体溯源报告']],
        widths=[0.5,2.4,1.8,1.3,1.5])

    add_heading(d, '数据来源说明', 1)
    add_table(d,
        ['数据项','来源系统/单位','获取方式','接口/路径说明','提供方','备注'],
        [['【数据项】','【系统名】','【API/库表/文件/人工】','【地址/表名/模板】','【方】',''],
         ['【数据项】','【系统名】','【方式】','【说明】','【方】',''],
         ['预警事件','驾驶舱后端（JSC）','数据库（库表）','warnings 表','我方',''],
         ['城运视频平台事件接入','全景影像视频平台（城运侧）→ 驾驶舱','HTTP 回调（入站）','POST /client/handle_event（事件）+ /client/handle_event_other（短视频）；平台按 V1.0 规范推送摄像头识别事件','城运中心（全景影像视频平台）','事件标准来源；eventId/cameraId/eventType/图片/区县街道'],
         ['城运事件推送','驾驶舱 → 三级城运中心','HTTP API（出站）','/api/smart-push/events → 城运中心事件接收接口；含自有 event_id','我方 + 城运中心','复用智治推送机制'],
         ['城运处置反馈','三级城运中心 → 驾驶舱','HTTP 回调（入站）','POST /api/smart-push/callback；带 X-Callback-Token 鉴权 + event_id 关联','城运中心','标准 schema + 字段映射可配'],
         ['处置与归档记录','驾驶舱后端（JSC）','数据库（库表）','warnings + disposition 扩展字段；结案PDF前端生成','我方',''],
         ['溯源数据（条件填充）','气体溯源系统 / 无人机气体巡查任务','关联任务数据（可选）','事件关联气体巡查任务时取热点/企业匹配/气象；无关联则标“不适用”','环保局/我方','版式参考，非强制依赖']],
        widths=[1.4,1.5,1.2,1.8,0.9,1.2])

    add_heading(d, '表格导入导出需求', 1)
    add_table(d,
        ['数据项','支持导入','支持导出','文件格式','字段说明','频率','操作权限','备注'],
        [['【数据项】','是/否','是/否','xlsx/csv','【字段清单】','【频率】','【角色】',''],
         ['处置记录','否','是','csv','处置措施/结果/责任单位/责任人/现场照片链接/政务派单号/处置时间/来源(城运/本地)','按需','管理员',''],
         ['城运推送与反馈记录','否','是','csv','event_id/城运事件号/推送状态/回调状态/推送时间/反馈时间/重试次数','按需','管理员','闭环留痕'],
         ['结案归档报告','否','是','PDF','完整结案报告（事件信息+证据+处置+溯源+结论）','按需（单条/批量）','管理员','前端 html2canvas+jsPDF 生成']],
        widths=[1.3,0.7,0.7,0.8,2.0,0.8,0.9,1.0])

    add_heading(d, '非功能性需求', 1)
    add_para(d, '性能：大屏需稳定承载约 30 路视频及实时数据展示，交互演示响应流畅（具体指标【请填写】）。')
    add_para(d, '可用性：支持向上级领导汇报场景的实景演示与交互操作；断网/单点故障时具备基本降级展示能力（【请填写】）。')
    add_para(d, '安全性：遵循电子政务网安全规范，数据访问按角色鉴权；城运回调接口采用共享令牌鉴权 + 来源 IP 白名单；处置/推送/回调/归档操作写入审计日志。')
    add_para(d, '兼容性：适配 2560×1440 及以上分辨率显示器及大屏投屏；浏览器/运行环境【请填写】。')

    # ===== 第九章（V1.2 含城运中心对接闭环）=====
    add_heading(d, '第九章 事件处置与结案归档模块详细需求（含三级城运中心对接闭环）', 1)
    add_para(d, '本章定义「预警 → 推送城运 → 政务处置 → 反馈回舱 → 结案归档」闭环的详细需求，作为开发实施的基准。')
    add_para(d, '设计原则：', bold=True)
    add_bullet(d, '复用驾驶舱现有预警体系（warnings 表）与既有“智治推送（城运中心对接）”机制（smart_push_events / smart_push_rules / smart_push_plans / smart_push_history）。')
    add_bullet(d, '新增“后半圈”：城运中心处置完成后通过回调接口将处置结果反馈回驾驶舱，自动更新事件状态并写回处置信息，形成完整执法闭环。')
    add_bullet(d, '未命中城运推送规则的预警，仍可由值守人员本地处置（handled），同样可生成结案报告，保证无对接盲区。')
    add_bullet(d, '结案报告版式严格对标万州区生态环境局《气体溯源系统》生成的“无人机大气巡查污染溯源报告”。')

    add_heading(d, '9.1 业务目标', 2)
    add_bullet(d, '对全部预警形成经三级城运中心流转处置的可追溯、可归档执法闭环记录。')
    add_bullet(d, '通过城运中心 API 自动推送事件、接收处置反馈，减少人工跨系统录入与状态跟办。')
    add_bullet(d, '一键生成风格统一、可直接归档/打印的 PDF 结案报告，减少人工整理文书的工作量。')
    add_bullet(d, '报告版式与既有气体溯源报告保持一致，便于环保局内部文书统一管理。')

    add_heading(d, '9.2 业务流程（闭环）', 2)
    add_para(d, '主链路（城运流转处置）：')
    add_bullet(d, '预警产生（warnings: pending）→ 命中城运推送规则（smart_push_rules）→ 驾驶舱经城运中心 API 推送事件（body 含自有 event_id 关联键）。')
    add_bullet(d, '驾驶舱将事件状态置为 pushed（已推送城运），并记录 smart_push_history。')
    add_bullet(d, '三级城运中心接收事件并流转至对应政务系统处置 → 处置中回调（status=processing）→ 驾驶舱置 processing。')
    add_bullet(d, '政务系统处置完成 → 城运中心回调（status=closed，带回处置措施/结果/责任单位/责任人/现场照片/政务派单号/处置时间/反馈时间）→ 驾驶舱写回 disposition、标记事件“待结案”（closed）。')
    add_bullet(d, '值守/管理员在驾驶舱对“待结案”事件一键生成 PDF 结案报告 → 报告纳入事件档案（archived=true）。')
    add_para(d, '兜底链路（本地处置）：', bold=True)
    add_bullet(d, '未命中推送规则的预警，由值守人员在驾驶舱本地填写处置信息（status: handled）， dispositionSource=local，同样可生成结案报告。')

    add_heading(d, '9.3 数据模型（后端扩展）', 2)
    add_para(d, '（1）warnings 状态机扩展（在 pending/handled 基础上新增城运闭环状态）：')
    add_table(d,
        ['状态','含义','触发方'],
        [['pending','预警产生，待处理','系统'],
         ['pushed','已推送城运中心，等待政务受理','驾驶舱推送成功'],
         ['processing','政务系统已受理、处置中','城运回调'],
         ['closed','政务处置完成，事件待结案','城运回调'],
         ['handled','本地手动处置完成（未推送城运的事件）','值守人员'],
         ['archived','已生成结案报告归档','值守/管理员']],
        widths=[1.2,3.6,1.8])
    add_para(d, '（2）处置信息扩展字段（建议存于 data_json 或独立 disposition 表；dispositionSource 区分来源）：')
    add_table(d,
        ['字段','类型','说明'],
        [['dispositionSource','枚举','处置来源：chengyun-callback（城运回调）/ local（本地录入）'],
         ['dispositionMeasures','文本','处置措施：现场复核 / 下达整改通知 / 移交执法 / 其他（可多选）'],
         ['dispositionResult','枚举','处置结果：已整改 / 持续观察 / 误报 / 移交执法'],
         ['responsibleUnit','文本','责任单位'],
         ['responsiblePerson','文本','责任人'],
         ['sitePhotos','URL 数组','现场照片 / 复核材料（可多张，城运回调或本地上传）'],
         ['govOrderNo','文本','政务派单号 / 处置单号（城运回调回传）'],
         ['dispositionNote','文本','备注（城运只读结论 + 本地可补充）'],
         ['dispositionTime','时间','处置完成时间'],
         ['feedbackTime','时间','城运回调反馈到达时间'],
         ['archived','布尔','是否已归档生成结案报告'],
         ['linkedTraceTaskId','文本','关联的气体巡查任务 ID（可选，用于条件填充溯源章节）']],
        widths=[1.8,1.0,4.0])
    add_para(d, '（3）城运事件映射与回调映射配置：')
    add_bullet(d, '事件映射：在 smart_push_history 或独立 event_mapping 表中存「我方 event_id ↔ 城运事件号（external_event_id）」，用于回调关联与去重。')
    add_bullet(d, 'callback_field_mapping：配置城运中心实际返回的 JSON 字段名 → 上述标准 disposition 字段的映射（与现有 push body_template 思路对称），适配不同政务系统接口。')

    add_heading(d, '9.4 处置录入与展示界面', 2)
    add_bullet(d, '城运回调自动处置：closed 回调到达后自动写回 disposition，界面以只读卡片展示“政务处置结果”（措施/结果/责任单位/责任人/现场照片/政务派单号），并允许值守人员补充备注。')
    add_bullet(d, '本地手动处置：未推送事件由值守人员填写处置表单（措施多选、结果下拉、责任单位、责任人、现场照片上传、备注）。')
    add_bullet(d, '待结案列表：status=closed 事件进入“待结案”视图，提供「生成结案报告」入口；status=pushed/processing 展示实时流转状态。')
    add_bullet(d, '证据联动：事件若含 AI 视频分析证据（图片/缩略图），自动带入结案报告证据区。')

    add_heading(d, '9.5 结案 PDF 模板（参照气体溯源报告版式，条件填充）', 2)
    add_para(d, '报告章节严格对标“无人机大气巡查污染溯源报告”，结构如下；标 ★ 的章节为溯源依赖章节，按事件类型条件填充。处置记录章内容取自 9.3 的 disposition 字段（来源可能是城运回调或本地录入）。')
    add_table(d,
        ['章节','内容','填充规则'],
        [['封面/标题','事件结案归档报告 · 事件编号/生成时间','固定'],
         ['一、事件概况','事件类型、级别、发生时间、地点、来源通道、关联企业（如有）、城运事件号','固定（取自预警数据 + govOrderNo）'],
         ['二、证据与监测分析 ★','AI 视频证据图、浓度/参数曲线','IoT视频事件填证据图；气体事件填监测分析（对标“污染监测分析”）'],
         ['三、气象条件分析 ★','主导风向、风速、扩散说明','仅关联气体巡查任务时填充；否则标“无溯源数据/不适用”'],
         ['四、溯源分析 ★','超标热点、嫌疑企业匹配（排名/距离/匹配分/特征污染物/风险等级）','仅关联气体巡查任务时填充；否则标“无溯源数据/不适用”'],
         ['五、处置记录','处置来源（城运/本地）、措施/结果/责任单位/责任人/现场照片/政务派单号/处置时间/反馈时间','固定（取自 disposition 字段）'],
         ['六、综合研判结论','事件结论、重点管控点位、后续建议','固定（可编辑）'],
         ['签章/归档信息','生成时间、技术支持、用途说明、归档编号','固定']],
        widths=[1.6,3.0,2.2])
    add_para(d, '说明：溯源章节（二/三/四中标★部分）完全复用气体溯源系统报告的表述逻辑与表格结构；当事件无关联气体巡查任务时，对应章节统一显示“无溯源数据 / 不适用”，不作任何虚构。')

    # ---- 9.5.1 示例：堆头未覆盖事件报告（基于系统真实配置填充）----
    add_heading(d, '9.5.1 事件报告示例——“堆头未覆盖”（基于驾驶舱真实推送规则填充）', 3)
    add_para(d, '以下以「堆头未覆盖」事件为例，按 9.5 模板结构给出一份完整填写的示例报告。数据取自驾驶舱系统「AI分析存档 → 推送规则 → 堆头未覆盖 24小时推送规则」的真实配置与 AI 视频分析存档（累计 38 条、九龙沙场 23 / 龙泗路 14 / GB_Chn_009 1）；标注※的为演示用占位值，非系统实测。')
    add_para(d, '该事件为 AI 视频分析类（堆头未覆盖），未关联无人机气体巡查任务，故第三章（气象）、第四章（溯源）按条件填充规则显示“无溯源数据 / 不适用”；第二章的“AI 视频证据”部分正常填充。', size=10, color=RGBColor(0x55,0x55,0x55))

    add_para(d, '【报告封面 / 标题】', bold=True)
    add_table(d,
        ['项目','内容'],
        [['报告名称','事件结案归档报告 · 堆头未覆盖（AI 视频分析）'],
         ['事件编号','JSC-20260710-DT-5633（※示例）'],
         ['生成时间','2026-07-11'],
         ['关联城运事件号','CY2026-0710-009812（※城运回调示例）'],
         ['触发推送规则','堆头未覆盖 24小时推送规则（阈值 1 / 时间窗 24h / 适用全部通道）']],
        widths=[2.0, 4.6])

    add_para(d, '一、事件概况', bold=True)
    add_table(d,
        ['字段','内容'],
        [['事件类型','堆头未覆盖（AI 视频分析识别）'],
         ['事件级别','2 级（一般 / 中度）'],
         ['发生时段（聚合窗）','2026-07-09 16:00 ～ 2026-07-10 15:58（24h 窗内累计）'],
         ['重点地点','九龙沙场'],
         ['来源通道','九龙沙场 NVR（通道号 56331706881318000004）'],
         ['聚合命中','24h 窗内同通道同类型累计识别 23 次（达阈值 1，触发城运推送）'],
         ['关联城运事件号','CY2026-0710-009812（※城运回调示例）']],
        widths=[2.0, 4.6])

    add_para(d, '二、证据与监测分析（★ 适用部分：AI 视频证据）', bold=True)
    add_para(d, '本事件为 IoT 视频 AI 识别类，第二章“证据与监测分析”中的 AI 视频证据部分正常填充；气体浓度监测曲线部分不适用，显示“无气体监测数据 / 不适用”。')
    add_bullet(d, '识别内容：物料堆场顶部未按扬尘管控要求苫盖防尘网（“堆头未覆盖”）。')
    add_bullet(d, '识别统计：24h 窗内累计 23 次，AI 置信度区间 0.51 ～ 0.60。')
    add_bullet(d, '证据图（样例 3 张，取自 AI 分析存档抓拍）：')
    add_table(d,
        ['序号','抓拍时间（上海时）','通道 / 地点','置信度','证据图地址（路径截断）'],
        [['1','2026-06-07 19:50','九龙沙场 NVR','0.51','…/images/detect/2026/06/07/2026-06-07T19:40:28.jpg'],
         ['2','2026-06-07 19:00','九龙沙场 NVR','0.54','…/images/detect/2026/06/07/2026-06-07T18:50:30.jpg'],
         ['3','2026-06-07 18:40','九龙沙场 NVR','0.60','…/images/detect/2026/06/07/2026-06-07T18:30:30.jpg']],
        widths=[0.6,1.7,1.3,0.9,2.1])
    add_para(d, '（说明：示例证据图地址取自系统 /images/detect/ 实际抓拍路径；前缀 http://111.10.220.226:5001/ 经图片代理 /api/iot-image 后可在驾驶舱内预览。）', size=9, color=RGBColor(0x55,0x55,0x55))

    add_para(d, '三、气象条件分析（★ 不适用）', bold=True)
    add_para(d, '无溯源数据 / 不适用：本事件为 AI 视频识别类（堆头未覆盖），未关联无人机气体巡查任务，无风向 / 风速 / 扩散等气象溯源数据，本章节显示“无溯源数据 / 不适用”，不作虚构。')

    add_para(d, '四、溯源分析（★ 不适用）', bold=True)
    add_para(d, '无溯源数据 / 不适用：本事件未关联气体巡查任务，无超标热点、嫌疑企业匹配（排名 / 距离 / 匹配分 / 特征污染物 / 风险等级）等溯源数据，本章节显示“无溯源数据 / 不适用”，不作虚构。')

    add_para(d, '五、处置记录', bold=True)
    add_para(d, '（以下为城运闭环示例：三级城运中心流转处置完成后，经回调接口 POST /api/smart-push/callback 将处置结果反馈回驾驶舱，自动写回 disposition；标注※为示例占位。）')
    add_table(d,
        ['字段','内容'],
        [['处置来源','城运回调（三级城运中心）'],
         ['处置措施','现场复核、下达整改通知（※示例）'],
         ['处置结果','已整改（※示例）'],
         ['责任单位','九龙沙场管理方（※示例）'],
         ['责任人','张××（※示例）'],
         ['现场照片','整改后苫盖照片 ×2（※示例链接）'],
         ['政务派单号','CY2026-0710-009812（※示例）'],
         ['处置时间','2026-07-10 17:30（※示例）'],
         ['反馈时间','2026-07-10 17:42（城运回调到达，※示例）']],
        widths=[2.0, 4.6])

    add_para(d, '六、综合研判结论', bold=True)
    add_bullet(d, '事件结论：九龙沙场物料堆场在统计时段内多次出现“堆头未覆盖”，存在扬尘污染风险；经三级城运中心流转处置，已完成防尘网苫盖整改。')
    add_bullet(d, '重点管控点位：九龙沙场堆场顶部区域。')
    add_bullet(d, '后续建议：建议加装堆场视频 AI 巡检与定期人工巡查，落实防尘网常态化苫盖，并纳入城运中心常态化监管。')

    add_para(d, '签章 / 归档信息', bold=True)
    add_table(d,
        ['项目','内容'],
        [['生成时间','2026-07-11'],
         ['技术支持','环保局驾驶舱系统（JSC）'],
         ['用途','环境执法结案归档'],
         ['归档编号','JSC-ARCH-20260711-0001（※示例）']],
        widths=[2.0, 4.6])

    add_heading(d, '9.6 权限与审计', 2)
    add_bullet(d, '值守人员：查看城运处置反馈、补充备注、对“待结案”事件生成结案报告；未推送事件可本地处置。')
    add_bullet(d, '管理员：配置城运推送规则/预案（smart_push_rules/plans）、回调字段映射（callback_field_mapping）、回调令牌与来源 IP 白名单、归档导出与 CSV 导出。')
    add_bullet(d, '所有推送（含重试）、回调接收、处置写回、结案生成操作均写入审计日志（操作人、时间、事件 ID / event_id），满足执法留痕与可审计要求。')

    add_heading(d, '9.7 三级城运中心对接（核心 · 含城运视频平台事件接入）', 2)
    add_para(d, '城运中心对接包含两条方向不同的链路，需分别实现：')
    add_bullet(d, '入站（城运视频平台 → 驾驶舱）：全景影像视频平台按《全景影像视频平台接口规范 V1.0》调用我方订阅接口推送摄像头识别事件与短视频。这是驾驶舱 AI 视频告警（如堆头未覆盖、扬尘等）的标准事件来源，对应 9.7.1 / 9.7.2。')
    add_bullet(d, '出站（驾驶舱 → 三级城运中心 / 政务系统）：预警经此链路推送至城运中心流转处置，处置结果经回调反馈回驾驶舱形成闭环。对应 9.7.3。')
    add_para(d, '基础地址（文档示例）：http://10.120.49.14:30100；实际 IP/端口由管理员提供（待确认）。在线 API 文档：http://10.120.49.14:30100/docs。', size=10, color=RGBColor(0x55,0x55,0x55))

    add_para(d, '9.7.1 事件接入（入站 · /client/handle_event）', bold=True)
    add_bullet(d, '接口：POST /client/handle_event（HTTP）。由我方按平台报文格式开发此订阅接口，全景影像视频平台调用该接口推送事件消息；响应固定返回 {"code":200,"message":"请求已成功","data":{}}。')
    add_bullet(d, '必填字段（平台 → 驾驶舱）：eventId（规则 sjzl-摄像头id-毫秒时间戳，如 sjzl-119-202409150714251）、eventTime、cameraId、eventType（枚举 1~17，见下方映射）、subType、elevation（上下角度）、azimuth（水平角度）、absoluteZoom（放大倍数）、eventImgSmall（小图网络地址）、eventImgBig（大图网络地址）、confirm（1已审核/0未审核）。')
    add_bullet(d, '选填字段：distance（预估距离）、latitude/longitude（经纬度；《全景影像视频平台接口规范 V1.0》原文对 latitude/longitude 含义描述存在互换笔误，已与城运确认：落库统一以「经度 = Longitude、纬度 = Latitude」为准，即取平台 longitude 字段落库为经度、latitude 字段落库为纬度）、address（位置）、watermarkImage（水印图）、count/total（识别物数量/总数）、presetPosNum/presetPosName（预置点）、processEventId/processEventStatus（流程事件标识/状态：0未识别、1识别到，可用于流程结束判断）、districtId/districtName（区县）、townId/townName（乡镇/街道）。')
    add_bullet(d, '落库映射：eventId→warning.id（保持平台事件ID，便于去重与回查）；cameraId + eventType→匹配 iot_channels 通道与 aiType；eventImgSmall/Big→picUrl（经 /api/iot-image 代理预览）；district/town/address→location；confirm→是否计入已审核告警。该接口推送的事件统一以 source=chengyun-platform 落库，与既有 iotcloud 来源并列。')
    add_para(d, '平台 eventType 枚举 → 驾驶舱 aiType 映射（已与城运确认）：', bold=True)
    add_table(d,
        ['平台 eventType','含义','驾驶舱 aiType（建议）'],
        [['1','工程车作业','工程车作业'],
         ['2','工程车数量','工程车数量'],
         ['3','烟尘','烟尘'],
         ['4','工地裸露地未覆盖','堆头未覆盖（已确认对应）'],
         ['5','生物质燃烧','生物质燃烧'],
         ['6','烟囱烟雾','烟囱烟雾'],
         ['7','扬尘','扬尘'],
         ['8','人员入侵','人员入侵'],
         ['9 / 10','卡车脏车 / 脏车','脏车'],
         ['12','建渣未覆盖','建渣未覆盖（独立类型，待确认是否纳入驾驶舱 aiType 体系）'],
         ['16','车辆冒装','车辆冒装'],
         ['17','工业烟羽','工业烟羽']],
        widths=[1.8,2.4,2.4])

    add_para(d, '9.7.2 短视频接入（入站 · /client/handle_event_other）', bold=True)
    add_bullet(d, '接口：POST /client/handle_event_other（HTTP）。平台推送事件关联短视频片段；参数 eventIds（事件ID）、fileUrl（录像文件地址）、cameraId。响应同 {"code":200,...}。')
    add_bullet(d, '落库：将 fileUrl 关联至对应 warning（按 eventIds 匹配），作为结案报告证据区的视频证据（优于静态图）。')

    add_para(d, '9.7.3 事件推送与处置反馈（出站 → 城运中心 / 政务系统闭环）', bold=True)
    add_bullet(d, '入口：POST /api/smart-push/events → checkRulesAndPush 按 smart_push_rules（event_type + 点位 + 触发阈值，复用既有智治推送机制）匹配 → 推送至三级城运中心事件接收接口。')
    add_bullet(d, '关联键：推送 body 通过 body_template 注入标准字段，并固定携带自有 event_id（本次推送聚合包 ID 或关联 warning id 列表），作为城运回调的关联键。')
    add_bullet(d, '推送失败处理：指数退避重试（最多 N 次），记录于 smart_push_history（request_body / response_status / 重试次数）。')
    add_bullet(d, '处置反馈回调（入站，新增）：POST /api/smart-push/callback（鉴权 + 关联 + 写回）；回传 event_id → 定位 warnings → 写回 disposition；按 event_id 去重。')
    add_bullet(d, '回调字段映射：按 callback_field_mapping 将城运返回字段映射到标准 disposition 字段（measures/result/responsibleUnit/responsiblePerson/photos/disposeTime/feedbackTime/govOrderNo），适配不同政务系统；标准回调 schema 与城运确认中（待提供规范）。')
    add_bullet(d, '状态驱动：城运回调 status=processing → 置 warnings 为 processing；status=closed → 置 closed + 写回 disposition + 标记“待结案”。')

    add_para(d, '9.7.4 可靠性与补偿', bold=True)
    add_bullet(d, '入站幂等：/client/handle_event 按 eventId 去重，重复推送不重复落库。')
    add_bullet(d, '出站重试：失败自动指数退避重试（最多 N 次，建议 3~5 次），记录于 smart_push_history。')
    add_bullet(d, '反馈超时告警：事件处于 pushed 状态超过 X 小时（建议 24h，待确认）仍无 processing/closed 反馈，标记为“反馈异常”并提示值守人员手动跟办，确保闭环不丢。')

    add_para(d, '9.7.5 安全与部署', bold=True)
    add_bullet(d, '入站鉴权（/client/handle_event、/client/handle_event_other）：已与城运确认，采用**政务网来源 IP 白名单**（平台出口 IP）作为主鉴权手段；**令牌鉴权功能预留**（接口支持请求头令牌校验，当前阶段不强制，待后续安全要求升级时启用）。V1.0 文档未规定接口鉴权，故以政务网隔离 + IP 白名单为落地方案。')
    add_bullet(d, '出站鉴权（/api/smart-push/callback）：X-Callback-Token（我方生成下发给城运中心）+ 政务网来源 IP 白名单，校验通过才受理；令牌不硬编码、定期轮换。')
    add_bullet(d, '接口部署于电子政务网隔离区，遵循政务网安全规范；入站与出站接口分别配置白名单与令牌。')
    add_para(d, '部署配置项（后端环境变量，经 systemd EnvironmentFile 注入，默认即可本地联调）：', size=10, color=RGBColor(0x55,0x55,0x55))
    add_bullet(d, 'CHENGYUN_ALLOW_IPS：入站接口允许的客户端 IP 白名单（逗号分隔）。默认 127.0.0.1（仅本机）；生产环境须设为城运视频平台出口 IP。不在白名单返回 403。')
    add_bullet(d, 'CHENGYUN_CALLBACK_TOKEN：可选。若设置，入站请求须带请求头 X-Callback-Token 且与之相等，否则返回 401。当前阶段不强制（令牌鉴权预留）。')
    add_bullet(d, 'CHENGYUN_IMG_HOSTS：城运平台图片/短视频域名白名单（逗号分隔），用于 /api/iot-image 代理放行。默认含文档示例 10.120.49.14；生产环境须设为平台真实图片域名。')

    add_heading(d, '验收标准', 1)
    add_bullet(d, '驾驶舱各功能模块按本说明书实现，与确认后的原型功能点一致。')
    add_bullet(d, '第四章数据清单所列指标均可正常展示，数据来源与第五章一致且实时/准时更新。')
    add_bullet(d, '第六章明确的导入导出功能可用，字段与权限符合要求。')
    add_bullet(d, '第九章事件处置归档模块可用：已处置预警可填写/接收处置信息、生成版式符合气体溯源报告的 PDF 结案报告；溯源章节按事件类型条件填充。')
    add_bullet(d, '城运中心对接可用：①入站事件接入——全景影像视频平台可按《V1.0 规范》调用我方 /client/handle_event（及 /client/handle_event_other 短视频）推送摄像头识别事件，字段正确落库为 warnings（source=chengyun-platform）且按 eventId 幂等去重；②出站闭环——预警可经 API 推送至城运中心（带 event_id）、政务处置结果可经回调接口（X-Callback-Token 鉴权）反馈回驾驶舱并自动写回处置、超时反馈有异常提示；推送与回调全程留痕。')
    add_bullet(d, '在约定硬件配置（见《需求沟通说明》）下，大屏展示与交互演示稳定流畅。')

    # ===== 附录二：接口契约明细（实现版 · 已落地）=====
    add_heading(d, '附录二：接口契约明细（实现版 · 已落地）', 1)
    add_para(d, '本章给出第九章两条对接链路的最终接口契约，与已部署后端代码严格一致，供城运中心集成联调与研发参考。所有响应体统一为 {"code":int,"message":str,"data":{...}} 结构。')
    add_para(d, 'A. 入站接口（城运视频平台 → 驾驶舱）', bold=True)

    add_para(d, 'A.1 /client/handle_event（事件推送）', bold=True)
    add_bullet(d, '方法/路径：POST /client/handle_event（经 nginx /client/ 转发至后端 7170）。')
    add_bullet(d, '鉴权：chengyunGuard（IP 白名单 CHENGYUN_ALLOW_IPS + 可选令牌 X-Callback-Token）；不在白名单返回 403，设了令牌但不符返回 401；/client/* 绕过 /api 会话鉴权。')
    add_table(d,
        ['字段','必填','说明'],
        [['eventId','是','事件ID（规则 sjzl-摄像头id-毫秒时间戳），作为 warning.id 幂等键'],
         ['eventTime','是','事件时间（上海本地时，无时区标记；落库补 +08:00）'],
         ['cameraId','是','摄像头ID → 关联 iot_channels 通道与通道名'],
         ['eventType','是','枚举 1~17 → aiType（映射见 9.7.1 表）'],
         ['confirm','否','1 已审核 / 0 未审核 → 决定 level（未审核=1，已审核=2）'],
         ['eventImgSmall / eventImgBig','否','小/大图网络地址 → picUrl（经 /api/iot-image 代理预览）'],
         ['latitude / longitude','否','经纬度；落库口径：经度=longitude、纬度=latitude（已与城运确认）'],
         ['districtName / townName / address','否','区县 / 街道 / 地址 → location'],
         ['elevation / azimuth / absoluteZoom','否','上下角度 / 水平角度 / 放大倍数'],
         ['processEventId / processEventStatus','否','流程事件标识 / 状态（0 未识别、1 识别到，用于流程结束判断）'],
         ['watermarkImage / count / total / presetPosNum / presetPosName / distance','否','水印图 / 识别物数量 / 总数 / 预置点编号 / 预置点名称 / 预估距离']],
        widths=[2.4, 0.6, 3.6])
    add_bullet(d, '响应：成功 200 {"code":200,"message":"请求已成功","data":{}}；缺 eventId 返回 400 {"code":400,"message":"缺少 eventId","data":{}}；非白名单 IP 返回 403；令牌不符返回 401。')
    add_bullet(d, '落库：以 eventId 为 warning.id 做幂等 upsert（INSERT OR REPLACE），重复推送不重复落库、保留首见时间与既有处置状态；source=chengyun-platform。')

    add_para(d, 'A.2 /client/handle_event_other（短视频接入）', bold=True)
    add_bullet(d, '方法/路径：POST /client/handle_event_other；鉴权与 A.1 相同。')
    add_bullet(d, '请求体：eventIds（事件ID，可逗号分隔）、fileUrl（录像文件地址）、cameraId。')
    add_bullet(d, '响应：200 {"code":200,"message":"请求已成功","data":{"updated":N}}（N 为成功关联的事件数）。')
    add_bullet(d, '行为：将 fileUrl 关联至对应 warning（按 eventIds 匹配），作为结案报告证据区的视频证据（优于静态图）。')

    add_para(d, 'B. 出站接口（驾驶舱 → 城运中心 / 政务系统）', bold=True)
    add_para(d, 'B.1 /api/smart-push/events（触发推送）', bold=True)
    add_bullet(d, '方法/路径：POST /api/smart-push/events；鉴权：会话鉴权（Bearer）。')
    add_bullet(d, '行为：按 smart_push_rules 匹配 → 调用预案 api_url 推送；推送成功关联事件置 pushed、历史置 pushed。')
    add_bullet(d, '推送报文模板变量（注入 body_template）：标准变量 event_type / location / lat / lon / level / value / standard / description / time / trigger_count / event_ids + 关联变量 push_id（本次 smart_push_history 记录 ID）+ callback_url（可选，配置 SMART_PUSH_CALLBACK_URL 时注入）。')
    add_bullet(d, '回执关联头：推送 HTTP 响应携带响应头 X-Push-Id: <push_id>，城运中心用于回调时关联本次推送。')

    add_para(d, 'B.2 /api/smart-push/callback（处置反馈回调 · 核心闭环）', bold=True)
    add_bullet(d, '方法/路径：POST /api/smart-push/callback；鉴权：复用 chengyunGuard（IP 白名单 CHENGYUN_ALLOW_IPS + 可选令牌 X-Callback-Token）。')
    add_bullet(d, '关联键（三选一，优先级从高到低）：请求头 X-Push-Id > 报文体 push_id > 报文体 event_id；三者皆缺返回 400 {"code":400,"message":"缺少 push_id（请携带 X-Push-Id 头或报文体 push_id）","data":{}}。')
    add_table(d,
        ['字段','必填','说明'],
        [['disposal_status','是','处置状态：processing（受理中）/ closed（已结案）；其他取值默认按 processing 处理'],
         ['disposal_result','否','处置结论 / 结果文本'],
         ['disposal_operator','否','处置人 / 责任单位'],
         ['disposal_time','否','处置完成时间'],
         ['(兼容别名)','否','status / result / remark / operator / disposalResult / disposalOperator / disposalTime，字段名可经 callback_field_mapping 适配不同政务系统']],
        widths=[2.4, 0.6, 3.6])
    add_bullet(d, '响应：成功 200 {"code":200,"message":"请求已成功","data":{"status":"<processing|closed>"}}；推送记录不存在 404 {"code":404,"message":"推送记录不存在","data":{}}；令牌不符 401；非白名单 IP 403。')
    add_bullet(d, '状态机（smart_push_events）：收到 processing → 关联事件置 processing（可回退）；收到 closed → 关联事件置 closed（不可逆，不得回退为 processing）+ 推送历史写回处置结论（disposal_result / disposal_operator / callback_time=反馈时间），标记"待结案"。')
    add_bullet(d, '超时判定：status=pushed 且 24h 内无任何回调，前端标记 is_timeout=1（红色告警），提示人工跟办，确保闭环不丢。')

    add_para(d, 'B.3 /api/smart-push/history/:id/close（人工一键结案）', bold=True)
    add_bullet(d, '方法/路径：POST /api/smart-push/history/:id/close；鉴权：会话鉴权（Bearer），操作人取当前登录用户。')
    add_bullet(d, '响应：成功 {"ok":true,"status":"closed"}；记录不存在 {"ok":false,"error":"推送记录不存在"}（HTTP 404）。')
    add_bullet(d, '行为：历史置 closed + 关联事件置 closed + 写回处置人 / 结案时间（disposal_result 默认"驾驶舱人工结案"）。')

    add_para(d, 'B.4 查询 /api/smart-push/history', bold=True)
    add_bullet(d, 'GET /api/smart-push/history?status=&event_type=&limit= ；status 支持 pushed / processing / closed / timeout（特殊：内存过滤 is_timeout=1）；event_type 按事件类型筛选；返回记录含 is_timeout 字段。')

    add_para(d, 'C. 部署配置项（/opt/jsc/backend/iotcloud.env，经 systemd EnvironmentFile 注入，改完 systemctl restart jsc-backend 生效）', bold=True)
    add_table(d,
        ['变量','默认值','说明'],
        [['CHENGYUN_ALLOW_IPS','127.0.0.1','入站 / 回调允许的客户端 IP 白名单（逗号分隔）；生产环境须设为城运视频平台出口 IP 与驾驶舱自身 IP'],
         ['CHENGYUN_CALLBACK_TOKEN','（空）','可选。若设置，入站 / 回调须带 X-Callback-Token 头且与之相等，否则 401；当前阶段不强制（令牌鉴权预留）'],
         ['CHENGYUN_IMG_HOSTS','10.120.49.14','城运平台图片 / 短视频域名白名单，供 /api/iot-image 代理放行'],
         ['SMART_PUSH_CALLBACK_URL','（空）','可选。驾驶舱对外回调地址，注入推送报文 callback_url 变量，便于城运中心回传']],
        widths=[2.2, 1.1, 3.3])

    add_heading(d, '附录：待环保局/城运中心确认问题清单', 1)
    add_table(d,
        ['序号','待确认问题','责任方','状态'],
        [['1','各展示数据的最终范围与口径','环保局','待确认'],
         ['2','各数据的来源系统、接口方式与提供方','环保局','待确认'],
         ['3','表格导入导出的数据项、字段与权限','环保局','待确认'],
         ['4','汇报演示场景下的降级与权限要求','环保局','待确认'],
         ['5','事件处置留痕字段是否满足执法存档（是否扩展/需电子签章）','环保局','待确认'],
         ['6','结案报告是否纳入统一档案管理/纸质归档/审计留存期限','环保局','待确认'],
         ['7','无溯源数据事件报告“溯源”章节标注“不适用”是否被接受','环保局','待确认'],
         ['8','处置与归档的操作权限划分（值守/管理员）','环保局','待确认'],
         ['9','城运中心事件接收接口地址、协议（HTTP/HTTPS）、字段规范','城运中心（三级城运）','待确认'],
         ['10','事件关联键 event_id 的命名/格式约定，以及城运是否回传其派单号','城运中心','待确认'],
         ['11','处置反馈回调接口规范：字段名、status 取值、推送时机（处置中/完成各一次或仅完成）','城运中心','待确认'],
         ['12','回调鉴权方式（X-Callback-Token 令牌交换、来源 IP 白名单范围）','城运中心','待确认'],
         ['13','推送重试次数 N 与反馈超时阈值 X（小时）的取值','城运中心/我方','待确认'],
         ['14','入站 /client/handle_event 的鉴权方式','城运中心（全景影像视频平台）','已确认：政务网来源 IP 白名单为主鉴权，令牌鉴权功能预留（当前不强制）'],
         ['15','平台 eventType 枚举（1~17）与驾驶舱 aiType 的精确映射','城运中心','已确认：“堆头未覆盖”对应平台 eventType=4（工地裸露地未覆盖）'],
         ['16','平台字段 latitude/longitude 经纬度落库口径','城运中心','已确认：落库以「经度=Longitude、纬度=Latitude」为准（latitude 字段落库为纬度、longitude 字段落库为经度）'],
         ['17','全景影像视频平台实际部署基础地址（文档示例 10.120.49.14:30100 仅为示例）','城运中心/管理员','待确认'],
         ['18','出站推送至三级城运中心 / 政务系统的接收接口地址与处置反馈回调规范（V1.0 文档未含，需另提供）','城运中心','待确认']],
        widths=[0.5,4.5,1.2,1.0])

    add_para(d, '编制单位：【请填写】          环保局（评审）：______________', size=10)
    add_para(d, '日期：2026年____月____日', size=10)

    d.save(path)
    print('saved doc2 ->', path)

if __name__ == '__main__':
    import os
    out_dir = r'E:\CC work\CC jsc\docs'
    os.makedirs(out_dir, exist_ok=True)
    build_doc1(os.path.join(out_dir, '环保局驾驶舱需求沟通说明_2026-07-11_含事件处置归档.docx'))
    build_doc2(os.path.join(out_dir, '环保局驾驶舱需求规格说明书_含事件处置归档_V1.6.docx'))
    print('DONE')
